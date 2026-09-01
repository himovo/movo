"""In-place DOCX translator.

Preserves original styling (fonts, colors, merged cells, headers/footers, images,
page layout, theme) by walking the python-docx object model and only replacing
text runs. The LLM never sees XML — it only receives plain text batches.

Design principles (see discussion doc):
- Glossary extraction (doc-level pass): ensures terminology consistency.
- Batch translation (8-10 segments per call): preserves intra-batch cohesion
  (pronouns, tense, transitions).
- Sliding context (neighbor segments): bridges batch boundaries.
- Length control via prompt: ±30% length, prefer concise phrasing.
- Column width auto-reflow for tables: avoids "narrow column holds long English"
  visual blow-up.
- Placeholders / numbers / dates / hyperlinks preserved verbatim.
"""

from __future__ import annotations

import io
import json
import logging
import re
from copy import deepcopy
from dataclasses import dataclass, field
from typing import Any, Dict, List

from docx import Document
from docx.oxml.ns import qn
from docx.section import _Footer
from docx.shared import Emu
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.llm.factory import get_llm_client
from app.llm.types import Message, Role


logger = logging.getLogger(__name__)


# ── Segment extraction ────────────────────────────────────────────────────────


@dataclass
class Segment:
    """One translatable text unit (a paragraph inside body/cell/header/footer)."""

    seg_id: int
    text: str
    kind: str  # "paragraph" | "cell" | "header" | "footer"
    paragraph_ref: Paragraph = field(repr=False)
    table_row_context: List[str] = field(default_factory=list)  # for cells
    table_column_header: str = ""  # for cells (best-effort)
    is_placeholder_only: bool = False  # no actual text to translate


_PLACEHOLDER_ONLY = re.compile(r"^\s*[\d\s\.\,\-\/\(\)\[\]\{\}#@%\*\+\$]*\s*$")


def _text_is_translatable(text: str) -> bool:
    """Skip empty, whitespace-only, or pure-number/symbol segments."""
    if not text or not text.strip():
        return False
    if _PLACEHOLDER_ONLY.match(text):
        return False
    return True


def _extract_row_context(row, skip_cell_idx: int) -> List[str]:
    """Collect short texts from other cells of the same row for context."""
    texts: List[str] = []
    for i, cell in enumerate(row.cells):
        if i == skip_cell_idx:
            continue
        t = cell.text.strip()
        if t and len(t) < 120:
            texts.append(t)
    return texts[:6]


def _iter_story_tables(container: Any) -> List[Table]:
    """Return all tables in a document/header/footer/cell, including nested ones."""
    out: List[Table] = []

    def visit_table(table: Table) -> None:
        out.append(table)
        for row in table.rows:
            seen_tc: set[int] = set()
            for cell in row.cells:
                tc_id = id(cell._element)
                if tc_id in seen_tc:
                    continue
                seen_tc.add(tc_id)
                for nested in list(getattr(cell, "tables", []) or []):
                    visit_table(nested)

    for table in list(getattr(container, "tables", []) or []):
        visit_table(table)
    return out


def _iter_header_footer_stories(doc: Document) -> List[Any]:
    stories: List[Any] = []
    seen: set[int] = set()
    story_refs = (
        ("header", "headerReference", "default"),
        ("first_page_header", "headerReference", "first"),
        ("even_page_header", "headerReference", "even"),
        ("footer", "footerReference", "default"),
        ("first_page_footer", "footerReference", "first"),
        ("even_page_footer", "footerReference", "even"),
    )
    for section in doc.sections:
        sect_pr = getattr(section, "_sectPr", None)
        for attr, ref_tag, ref_type in story_refs:
            if sect_pr is not None:
                refs = sect_pr.findall(qn(f"w:{ref_tag}"))
                if not any((ref.get(qn("w:type")) or "default") == ref_type for ref in refs):
                    continue
            try:
                story = getattr(section, attr)
            except Exception:
                continue
            story_id = id(getattr(story, "_element", story))
            if story_id in seen:
                continue
            seen.add(story_id)
            stories.append(story)
    return stories


def _extract_table_segments(
    table: Table,
    *,
    add_paragraph: Any,
    seen_paragraph_ids: set[int],
) -> None:
    """Extract paragraphs from one table, deduping merged-cell aliases."""
    header_texts: List[str] = []
    try:
        header_row = table.rows[0]
        seen_header_tc: set[int] = set()
        for c in header_row.cells:
            tc_id = id(c._element)
            if tc_id in seen_header_tc:
                header_texts.append("")
                continue
            seen_header_tc.add(tc_id)
            header_texts.append(c.text.strip())
    except Exception:
        header_texts = []

    for row_idx, row in enumerate(table.rows):
        seen_tc_in_row: set[int] = set()
        row_ctx_cells: List[str] = []
        for c in row.cells:
            tc_id = id(c._element)
            if tc_id in seen_tc_in_row:
                continue
            seen_tc_in_row.add(tc_id)
            t = c.text.strip()
            if t:
                row_ctx_cells.append(t)
        seen_tc_in_row.clear()
        for cell_idx, cell in enumerate(row.cells):
            tc_id = id(cell._element)
            if tc_id in seen_tc_in_row:
                continue
            seen_tc_in_row.add(tc_id)
            column_header = ""
            if row_idx > 0 and cell_idx < len(header_texts):
                column_header = header_texts[cell_idx]
            for p in cell.paragraphs:
                p_id = id(p._p)
                if p_id in seen_paragraph_ids:
                    continue
                if not _text_is_translatable(p.text):
                    continue
                seen_paragraph_ids.add(p_id)
                add_paragraph(
                    p,
                    "cell",
                    row_context=[t for t in row_ctx_cells if t != p.text.strip()][:6],
                    column_header=column_header,
                )


def extract_segments(doc: Document) -> List[Segment]:
    """Walk the document, produce an ordered list of translatable segments."""
    segments: List[Segment] = []
    next_id = 0

    def _add_paragraph(p: Paragraph, kind: str, **meta) -> None:
        nonlocal next_id
        txt = p.text
        if not _text_is_translatable(txt):
            return
        segments.append(
            Segment(
                seg_id=next_id,
                text=txt,
                kind=kind,
                paragraph_ref=p,
                table_row_context=list(meta.get("row_context") or []),
                table_column_header=str(meta.get("column_header") or ""),
            )
        )
        next_id += 1

    # Body paragraphs (non-table top-level)
    for p in doc.paragraphs:
        _add_paragraph(p, "paragraph")

    # Tables (including nested): header row provides column context.
    # Merged cells expose the same underlying <w:tc> element multiple times — we dedupe
    # by element identity so merged ranges are translated once, not N times.
    seen_paragraph_ids: set[int] = set()
    for table in _iter_story_tables(doc):
        _extract_table_segments(table, add_paragraph=_add_paragraph, seen_paragraph_ids=seen_paragraph_ids)

    # Headers and footers of each section
    for story in _iter_header_footer_stories(doc):
        kind = "footer" if isinstance(story, _Footer) else "header"
        try:
            for p in story.paragraphs:
                _add_paragraph(p, kind)
            for table in _iter_story_tables(story):
                _extract_table_segments(table, add_paragraph=_add_paragraph, seen_paragraph_ids=seen_paragraph_ids)
        except Exception:
            pass

    return segments


# ── Glossary extraction ───────────────────────────────────────────────────────


async def extract_glossary(
    *,
    llm: Any,
    full_text_sample: str,
    source_language: str,
    target_language: str,
) -> Dict[str, str]:
    """Ask LLM to build a term consistency table for the document.

    Returns a dict {source_term: target_term}. On failure returns an empty dict.
    """
    system = (
        "You extract a minimal glossary of terms that MUST be translated consistently.\n"
        f"Translate from {source_language} to {target_language}.\n"
        "Include: proper nouns (companies, products, people), legal/contract terms, "
        "defined abbreviations, recurring domain-specific terms.\n"
        "EXCLUDE: generic words, pronouns, single-occurrence words.\n"
        'Return strict JSON: {"glossary": [{"source": "...", "target": "..."}, ...]}\n'
        "Keep it under 40 entries."
    )
    payload = {"document_excerpt": full_text_sample[:8000]}
    try:
        resp = await llm.ainvoke(
            [
                Message(role=Role.SYSTEM, content=system),
                Message(role=Role.USER, content=json.dumps(payload, ensure_ascii=False)),
            ]
        )
        raw = str(getattr(resp, "content", "") or "").strip()
        cleaned = raw.replace("```json", "").replace("```", "").strip()
        data = json.loads(cleaned)
        out: Dict[str, str] = {}
        for item in (data or {}).get("glossary") or []:
            if not isinstance(item, dict):
                continue
            src = str(item.get("source") or "").strip()
            tgt = str(item.get("target") or "").strip()
            if src and tgt and src != tgt:
                out[src] = tgt
        return out
    except Exception as exc:
        logger.warning("[docx_inplace] glossary_extraction_failed: %s", exc)
        return {}


# ── Batch translation ─────────────────────────────────────────────────────────


_TRANSLATION_SYSTEM = """You are a professional document translator.
Translate from {source_language} to {target_language}.

Rules:
- Preserve ALL numbers, dates, currency, placeholders (e.g. #{{name}}, {{var}}), URLs verbatim
- Keep length CLOSE to source (within ±30% characters); prefer concise phrasing
- Keep proper nouns, brand names, technical terms as-is unless the glossary says otherwise
- Use the provided glossary strictly for listed terms
- Translate EACH segment independently to its target; do not merge, split, or reorder segments
- For segments marked "kind":"cell", prefer short concise phrasing (table cells are narrow)
- For segments marked "kind":"header" / "footer", keep very concise
- Do NOT add translator notes, explanations, or extra punctuation
- Output strict JSON with the same segment ids
"""


async def translate_batch(
    *,
    llm: Any,
    batch: List[Segment],
    source_language: str,
    target_language: str,
    glossary: Dict[str, str],
    context_before: str = "",
    context_after: str = "",
) -> Dict[int, str]:
    """Translate a batch of segments. Returns {seg_id: translated_text}."""
    if not batch:
        return {}
    system = _TRANSLATION_SYSTEM.format(
        source_language=source_language,
        target_language=target_language,
    )
    payload: Dict[str, Any] = {
        "glossary": [{"source": k, "target": v} for k, v in glossary.items()][:40],
        "context_before": context_before[:400],
        "context_after": context_after[:400],
        "segments": [
            {
                "id": seg.seg_id,
                "kind": seg.kind,
                "text": seg.text,
                **({"row_context": seg.table_row_context} if seg.table_row_context else {}),
                **({"column_header": seg.table_column_header} if seg.table_column_header else {}),
            }
            for seg in batch
        ],
    }
    user_instruction = (
        "Translate each segment. Return strict JSON:\n"
        '{"translations":[{"id":<int>,"text":"<translated>"}]}'
    )
    try:
        resp = await llm.ainvoke(
            [
                Message(role=Role.SYSTEM, content=system),
                Message(role=Role.USER, content=user_instruction + "\n\n" + json.dumps(payload, ensure_ascii=False)),
            ]
        )
        raw = str(getattr(resp, "content", "") or "").strip()
        cleaned = raw.replace("```json", "").replace("```", "").strip()
        data = json.loads(cleaned)
        out: Dict[int, str] = {}
        for item in (data or {}).get("translations") or []:
            if not isinstance(item, dict):
                continue
            try:
                sid = int(item.get("id"))
            except Exception:
                continue
            txt = str(item.get("text") or "").strip()
            if txt:
                out[sid] = txt
        return out
    except Exception as exc:
        logger.warning("[docx_inplace] batch_translation_failed: %s", exc)
        return {}


# ── Write-back ────────────────────────────────────────────────────────────────


def _copy_paragraph_run_defaults(paragraph: Paragraph, run: Any) -> None:
    ppr = getattr(getattr(paragraph, "_p", None), "pPr", None)
    source_rpr = ppr.find(qn("w:rPr")) if ppr is not None else None
    if source_rpr is None:
        return
    r = getattr(run, "_r", None)
    if r is None:
        return
    existing_rpr = getattr(r, "rPr", None)
    if existing_rpr is not None:
        r.remove(existing_rpr)
    r.insert(0, deepcopy(source_rpr))


def write_translation_into_paragraph(p: Paragraph, translated: str) -> None:
    """Replace paragraph text, preserving the style of its first run.

    Strategy: keep the first run's formatting, put the entire translated text
    into its .text, clear all other runs. This loses intra-paragraph run-level
    style variation (e.g. one bold word mid-sentence) but reliably keeps the
    paragraph's overall font, size, color, alignment, indentation, etc.
    """
    runs = list(p.runs)
    if not runs:
        run = p.add_run(translated)
        _copy_paragraph_run_defaults(p, run)
        return
    runs[0].text = translated
    for r in runs[1:]:
        r.text = ""


# ── Table column width reflow ────────────────────────────────────────────────


def _cell_longest_line_chars(cell) -> int:
    longest = 0
    for p in cell.paragraphs:
        for line in (p.text or "").splitlines() or [p.text or ""]:
            longest = max(longest, len(line))
    return longest


def reflow_table_column_widths(doc: Document) -> int:
    """Rebalance column widths based on longest translated text per column.

    Keeps each table's total width constant. Operates on non-merged column
    boundaries; merged cells (gridSpan/vMerge) are unaffected because we only
    set individual column preferred widths.

    Returns the number of tables whose widths were adjusted.
    """
    adjusted = 0
    tables = list(_iter_story_tables(doc))
    for story in _iter_header_footer_stories(doc):
        tables.extend(_iter_story_tables(story))
    seen_tables: set[int] = set()
    for table in tables:
        table_id = id(table._element)
        if table_id in seen_tables:
            continue
        seen_tables.add(table_id)
        try:
            rows = list(table.rows)
            if not rows:
                continue
            num_cols = max(len(r.cells) for r in rows)
            if num_cols < 2:
                continue
            # Measure each column's longest character count
            col_weights = [1] * num_cols
            for r in rows:
                for i, cell in enumerate(r.cells):
                    if i >= num_cols:
                        break
                    col_weights[i] = max(col_weights[i], _cell_longest_line_chars(cell))
            # Skip tables where column weights are already balanced
            total_weight = sum(col_weights)
            if total_weight <= num_cols:
                continue
            # Collect existing column widths
            existing_widths = []
            for col in table.columns:
                try:
                    existing_widths.append(int(col.width) if col.width is not None else 0)
                except Exception:
                    existing_widths.append(0)
            total_width = sum(existing_widths)
            if total_width <= 0:
                continue
            # Blend existing ratio (0.4) with new weight-based ratio (0.6)
            new_widths: List[int] = []
            for i in range(num_cols):
                existing_ratio = existing_widths[i] / total_width if i < len(existing_widths) else 1 / num_cols
                new_ratio = col_weights[i] / total_weight
                blended = 0.4 * existing_ratio + 0.6 * new_ratio
                new_widths.append(int(total_width * blended))
            # Normalize so the sum equals total_width exactly
            delta = total_width - sum(new_widths)
            if new_widths:
                new_widths[-1] += delta
            # Apply
            for i, col in enumerate(table.columns):
                if i >= len(new_widths):
                    break
                try:
                    col.width = Emu(max(1, new_widths[i]))
                except Exception:
                    pass
            adjusted += 1
        except Exception as exc:
            logger.debug("[docx_inplace] table_reflow_skip: %s", exc)
    return adjusted


# ── Orchestrator ──────────────────────────────────────────────────────────────


@dataclass
class TranslationResult:
    file_bytes: bytes
    stats: Dict[str, Any]
    glossary: Dict[str, str]


async def translate_docx_inplace(
    *,
    source_bytes: bytes,
    source_language: str,
    target_language: str,
    llm: Any = None,
    batch_size: int = 8,
) -> TranslationResult:
    """Translate a .docx file in place, preserving all styling.

    Parameters
    ----------
    source_bytes : raw bytes of the source .docx file
    source_language / target_language : language names (e.g. "Chinese", "English")
    llm : optional LLM client (defaults to the "compose" stage client)
    batch_size : number of segments per LLM call
    """
    if llm is None:
        llm = get_llm_client(streaming=False, stage="compose")

    doc = Document(io.BytesIO(source_bytes))
    segments = extract_segments(doc)
    if not segments:
        buf = io.BytesIO()
        doc.save(buf)
        return TranslationResult(
            file_bytes=buf.getvalue(),
            stats={"segments": 0, "batches": 0, "tables_reflowed": 0},
            glossary={},
        )

    # 1) Glossary pass — sample up to ~8k chars of document text for LLM
    full_sample = "\n".join(seg.text for seg in segments)[:8000]
    glossary = await extract_glossary(
        llm=llm,
        full_text_sample=full_sample,
        source_language=source_language,
        target_language=target_language,
    )

    # 2) Batch translation with sliding context
    translations: Dict[int, str] = {}
    total = len(segments)
    num_batches = (total + batch_size - 1) // batch_size
    for b_idx in range(num_batches):
        start = b_idx * batch_size
        end = min(start + batch_size, total)
        batch = segments[start:end]
        context_before = ""
        context_after = ""
        if start > 0:
            context_before = segments[start - 1].text
        if end < total:
            context_after = segments[end].text
        batch_result = await translate_batch(
            llm=llm,
            batch=batch,
            source_language=source_language,
            target_language=target_language,
            glossary=glossary,
            context_before=context_before,
            context_after=context_after,
        )
        translations.update(batch_result)

    # 3) Write back in place
    written = 0
    for seg in segments:
        translated = translations.get(seg.seg_id)
        if not translated:
            continue  # leave original text if translation missing
        write_translation_into_paragraph(seg.paragraph_ref, translated)
        written += 1

    # 4) Reflow table column widths so long translations don't crush narrow columns
    tables_reflowed = reflow_table_column_widths(doc)

    # 5) Save
    buf = io.BytesIO()
    doc.save(buf)
    return TranslationResult(
        file_bytes=buf.getvalue(),
        stats={
            "segments": total,
            "segments_translated": written,
            "batches": num_batches,
            "tables_reflowed": tables_reflowed,
            "glossary_size": len(glossary),
        },
        glossary=glossary,
    )


__all__ = [
    "Segment",
    "TranslationResult",
    "extract_segments",
    "extract_glossary",
    "translate_batch",
    "write_translation_into_paragraph",
    "reflow_table_column_widths",
    "translate_docx_inplace",
]
