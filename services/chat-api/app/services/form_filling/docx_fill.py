from __future__ import annotations

import io
import re
from copy import deepcopy
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.oxml.ns import qn

from app.services.form_filling.mapper import build_fill_plan
from app.services.form_filling.models import FillResult


def _paragraphs_in_cell(cell: Any) -> Iterable[Any]:
    for paragraph in list(getattr(cell, "paragraphs", []) or []):
        yield paragraph
    for table in list(getattr(cell, "tables", []) or []):
        for row in table.rows:
            for nested_cell in row.cells:
                yield from _paragraphs_in_cell(nested_cell)


def _all_paragraphs(doc: Any) -> Iterable[Any]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _paragraphs_in_cell(cell)


def _copy_paragraph_run_defaults(paragraph: Any, run: Any) -> None:
    ppr = getattr(getattr(paragraph, "_p", None), "pPr", None)
    source_rpr = ppr.find(qn("w:rPr")) if ppr is not None else None
    if source_rpr is None:
        return
    r = getattr(run, "_r", None)
    if r is None:
        return
    existing_rpr = getattr(r, "rPr", None)
    if existing_rpr is not None:
        r.remove(existing_rpr)
    r.insert(0, deepcopy(source_rpr))


def _replace_paragraph_text(paragraph: Any, old: str, new: str) -> bool:
    text = paragraph.text or ""
    if old not in text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if paragraph.runs:
        paragraph.runs[0].text = text.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""
        return True
    return False


def _set_cell_text(cell: Any, value: str) -> None:
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
            return
        run = paragraph.add_run(value)
        _copy_paragraph_run_defaults(paragraph, run)
        return
    cell.add_paragraph(value)


def _cell_text(cell: Any) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _logical_row_cells(row: Any) -> List[Dict[str, Any]]:
    """Return visual/logical cells, collapsing python-docx merged-cell repeats."""
    out: List[Dict[str, Any]] = []
    last_tc_id: int | None = None
    for idx, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id == last_tc_id:
            continue
        out.append({"index": idx, "cell": cell, "text": _cell_text(cell)})
        last_tc_id = tc_id
    return out


def _non_empty_texts(items: List[Dict[str, Any]]) -> List[str]:
    return [str(item.get("text") or "").strip() for item in items if str(item.get("text") or "").strip()]


def extract_docx_form_schema(source_bytes: bytes) -> Dict[str, Any]:
    doc = Document(io.BytesIO(source_bytes))
    fill_targets: List[Dict[str, Any]] = []
    tables: List[Dict[str, Any]] = []
    row_groups: List[Dict[str, Any]] = []

    for p_idx, paragraph in enumerate(_all_paragraphs(doc)):
        text = paragraph.text or ""
        for name in re.findall(r"\{\{\s*([^{}]{1,50})\s*\}\}", text):
            fill_targets.append(
                {
                    "target": f"paragraph:{p_idx}:placeholder:{name.strip()}",
                    "label": name.strip(),
                    "kind": "placeholder",
                    "placeholder": "{{" + name + "}}",
                }
            )

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            logical_cells = _logical_row_cells(row)
            non_empty = _non_empty_texts(logical_cells)
            if len(non_empty) >= 2:
                headers = [
                    {
                        "header": str(item.get("text") or "").strip(),
                        "cell_col": int(item.get("index") or 0),
                    }
                    for item in logical_cells
                    if str(item.get("text") or "").strip()
                ][:30]
                tables.append(
                    {
                        "target": f"table:{t_idx}:row:{r_idx}",
                        "table_index": t_idx,
                        "header_row": r_idx,
                        "headers": [h["header"] for h in headers],
                        "header_cells": headers,
                    }
                )
                if len(headers) >= 4 and r_idx + 1 < len(table.rows):
                    data_rows: List[int] = []
                    section = str(headers[0].get("header") or "").strip()
                    data_columns = [h for h in headers[1:] if str(h.get("header") or "").strip()]
                    for next_idx in range(r_idx + 1, len(table.rows)):
                        next_cells = _logical_row_cells(table.rows[next_idx])
                        next_texts = _non_empty_texts(next_cells)
                        if len(next_texts) >= 2 and next_idx != r_idx + 1:
                            break
                        first_text = str(next_cells[0].get("text") or "").strip() if next_cells else ""
                        if section and first_text and first_text != section and next_texts:
                            break
                        if any(not _cell_text(table.rows[next_idx].cells[int(col.get("cell_col") or 0)]) for col in data_columns):
                            data_rows.append(next_idx)
                    if data_rows and data_columns:
                        row_groups.append(
                            {
                                "target": f"table:{t_idx}:row_group:{r_idx}",
                                "table_index": t_idx,
                                "header_row": r_idx,
                                "section": section,
                                "data_rows": data_rows[:20],
                                "columns": data_columns[:20],
                            }
                        )
            previous_first_text = ""
            if r_idx > 0:
                previous_cells = _logical_row_cells(table.rows[r_idx - 1])
                previous_first_text = str(previous_cells[0].get("text") or "").strip() if previous_cells else ""
            for pos, item in enumerate(logical_cells):
                c_idx = int(item.get("index") or 0)
                text = str(item.get("text") or "")
                if not text:
                    continue
                repeated_section_label = pos == 0 and len(non_empty) == 1 and previous_first_text and text.strip() == previous_first_text
                if (
                    not repeated_section_label
                    and pos + 1 < len(logical_cells)
                    and not str(logical_cells[pos + 1].get("text") or "").strip()
                ):
                    target_idx = int(logical_cells[pos + 1].get("index") or (c_idx + 1))
                    fill_targets.append(
                        {
                            "target": f"table:{t_idx}:cell:{r_idx}:{target_idx}",
                            "label": text[:80],
                            "kind": "right_empty_cell",
                        }
                    )
                if r_idx + 1 < len(table.rows) and not _cell_text(table.rows[r_idx + 1].cells[c_idx]):
                    fill_targets.append(
                        {
                            "target": f"table:{t_idx}:cell:{r_idx + 1}:{c_idx}",
                            "label": text[:80],
                            "kind": "below_empty_cell",
                        }
                    )
    return {"fill_targets": fill_targets[:300], "tables": tables[:80], "row_groups": row_groups[:40]}


def _append_docx_row(table: Any, source_row_idx: int) -> Any:
    source_row = table.rows[source_row_idx]
    new_tr = deepcopy(source_row._tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for cell in row.cells:
        _set_cell_text(cell, "")
    return row


async def fill_docx_form(
    *,
    source_bytes: bytes,
    user_text: str,
    llm: Any,
    overwrite: bool = False,
) -> FillResult:
    schema = extract_docx_form_schema(source_bytes)
    plan = await build_fill_plan(llm=llm, user_text=user_text, schema=schema, file_type="docx", overwrite=overwrite)
    doc = Document(io.BytesIO(source_bytes))
    warnings: List[str] = [str(w) for w in list(plan.get("warnings") or []) if str(w).strip()]
    filled = 0
    appended = 0
    skipped = 0

    paragraphs = list(_all_paragraphs(doc))
    for item in list(plan.get("fills") or []):
        if not isinstance(item, dict):
            continue
        target = str(item.get("target") or "")
        value = str(item.get("value") or "").strip()
        if not value:
            continue
        if target.startswith("paragraph:"):
            parts = target.split(":")
            try:
                paragraph = paragraphs[int(parts[1])]
            except Exception:
                skipped += 1
                continue
            placeholder = str(item.get("placeholder") or "")
            label = str(item.get("label") or "").strip()
            candidates = [placeholder, "{{" + label + "}}", "{{ " + label + " }}"]
            if any(_replace_paragraph_text(paragraph, token, value) for token in candidates if token):
                filled += 1
            else:
                skipped += 1
            continue
        if target.startswith("table:"):
            parts = target.split(":")
            try:
                table = doc.tables[int(parts[1])]
                cell = table.rows[int(parts[3])].cells[int(parts[4])]
            except Exception:
                skipped += 1
                continue
            if _cell_text(cell) and not overwrite:
                skipped += 1
                continue
            _set_cell_text(cell, value)
            filled += 1

    row_group_index = {str(t.get("target") or ""): t for t in list(schema.get("row_groups") or []) if isinstance(t, dict)}
    for item in list(plan.get("fill_rows") or []):
        if not isinstance(item, dict):
            continue
        group_meta = row_group_index.get(str(item.get("target") or ""))
        if not group_meta:
            continue
        rows_payload = item.get("rows") if isinstance(item.get("rows"), list) else []
        if not rows_payload:
            values = item.get("values") if isinstance(item.get("values"), dict) else {}
            rows_payload = [values] if values else []
        if not rows_payload:
            continue
        try:
            table = doc.tables[int(group_meta.get("table_index") or 0)]
            data_rows = [int(x) for x in list(group_meta.get("data_rows") or [])]
            columns = [dict(x) for x in list(group_meta.get("columns") or []) if isinstance(x, dict)]
        except Exception:
            skipped += 1
            continue
        for row_values, row_idx in zip(rows_payload, data_rows):
            if not isinstance(row_values, dict):
                continue
            for col in columns:
                header = str(col.get("header") or "").strip()
                if header not in row_values:
                    continue
                value = str(row_values.get(header) or "").strip()
                if not value:
                    continue
                try:
                    cell = table.rows[row_idx].cells[int(col.get("cell_col") or 0)]
                except Exception:
                    skipped += 1
                    continue
                if _cell_text(cell) and not overwrite:
                    continue
                _set_cell_text(cell, value)
                filled += 1

    table_index = {str(t.get("target") or ""): t for t in list(schema.get("tables") or []) if isinstance(t, dict)}
    for item in list(plan.get("append_rows") or []):
        if not isinstance(item, dict):
            continue
        table_meta = table_index.get(str(item.get("target") or ""))
        values = dict(item.get("values") or {}) if isinstance(item.get("values"), dict) else {}
        if not table_meta or not values:
            continue
        try:
            table = doc.tables[int(table_meta.get("table_index") or 0)]
            header_row = int(table_meta.get("header_row") or 0)
            headers = [str(h or "").strip() for h in list(table_meta.get("headers") or [])]
            header_cells = [
                dict(x)
                for x in list(table_meta.get("header_cells") or [])
                if isinstance(x, dict) and str(x.get("header") or "").strip()
            ]
            row = _append_docx_row(table, min(header_row + 1, len(table.rows) - 1))
        except Exception:
            skipped += 1
            continue
        if header_cells:
            for meta in header_cells:
                header = str(meta.get("header") or "").strip()
                col_idx = int(meta.get("cell_col") or 0)
                if header in values and col_idx < len(row.cells):
                    _set_cell_text(row.cells[col_idx], str(values[header]))
        else:
            for idx, header in enumerate(headers):
                if idx < len(row.cells) and header in values:
                    _set_cell_text(row.cells[idx], str(values[header]))
        appended += 1

    cleared_placeholders = 0
    unresolved: List[str] = []
    for paragraph in _all_paragraphs(doc):
        text = paragraph.text or ""
        for name in re.findall(r"\{\{\s*([^{}]{1,50})\s*\}\}", text):
            label = str(name or "").strip()
            token_variants = ["{{" + label + "}}", "{{ " + label + " }}"]
            changed = False
            for token in token_variants:
                if _replace_paragraph_text(paragraph, token, ""):
                    changed = True
            if changed:
                cleared_placeholders += 1
                if label:
                    unresolved.append(label)

    out = io.BytesIO()
    doc.save(out)
    return FillResult(
        file_bytes=out.getvalue(),
        stats={
            "filled_fields": filled,
            "appended_rows": appended,
            "skipped": skipped,
            "cleared_placeholders": cleared_placeholders,
            "planned_fills": len(list(plan.get("fills") or [])),
            "planned_fill_rows": len(list(plan.get("fill_rows") or [])),
            "planned_append_rows": len(list(plan.get("append_rows") or [])),
        },
        warnings=warnings + [f"No value provided for '{label}'; placeholder cleared." for label in unresolved[:20]],
    )
