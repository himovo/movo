from __future__ import annotations
from app.infrastructure.observability.config import log_print

import os
import mimetypes
import re
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from docx.text.paragraph import Paragraph
from app.tools.tooling import tool

from app.core.config import get_settings
from app.utils.i18n import get as i18n_get
from app.utils.i18n import get_list as i18n_get_list
from app.utils.oss_uploader import AliyunOSSUploader


IMAGE_RE = re.compile(r"^!\[.*?\]\((.*?)\)")
HTML_IMG_RE = re.compile(r'<img[^>]+src=["\']([^"\']+)["\']', re.IGNORECASE)


def _sanitize_image_caption(value: str) -> str:
    caption = str(value or "").strip()
    caption = re.sub(r"[\r\n\t]+", " ", caption)
    caption = re.sub(r"\s{2,}", " ", caption).strip()
    caption = caption.replace("[", "（").replace("]", "）")
    return caption[:160] or "image"


def _normalize_image_markers(markdown: str) -> str:
    text = str(markdown or "")
    if not text:
        return text

    def repl(match: re.Match[str]) -> str:
        caption = _sanitize_image_caption(match.group(1))
        target = str(match.group(2) or "").strip()
        if target.startswith("<") and target.endswith(">"):
            target = target[1:-1].strip()
        return f"![{caption}]({target})"

    text = re.sub(r"!\[(.*?)\]\(<([^>]+)>\)", repl, text, flags=re.DOTALL)
    text = re.sub(r"!\[(.*?)\]\(([^)\r\n]+)\)", repl, text, flags=re.DOTALL)
    return text


def _extract_image_url(raw: str) -> str:
    """
    Extract the URL from a markdown image target, handling optional title and <...> wrappers.
    """
    if not raw:
        return ""
    target = raw.strip()
    if target.startswith("<") and target.endswith(">"):
        target = target[1:-1].strip()
    # Remove optional title: url "title" or url 'title'
    if " " in target:
        parts = target.split()
        if parts:
            target = parts[0].strip()
    return target


def _resolve_render_image_url(value: str) -> str:
    url = str(value or "").strip()
    if not url or url.startswith(("http://", "https://", "data:", "file://")):
        return url
    if url.startswith("/"):
        base = str(get_settings().BACKEND_INTERNAL_BASE_URL or "http://127.0.0.1:8000").strip().rstrip("/")
        public_prefix = "/" + str(get_settings().FILE_PUBLIC_PATH_PREFIX or "/askai-api/api/files").strip().strip("/")
        if url.startswith(public_prefix + "/"):
            url = "/api/files/" + url[len(public_prefix) + 1 :]
        return f"{base}{url}"
    return url


def _guess_image_suffix(*, url: str, object_path: str = "", content_type: str = "") -> str:
    guessed = ""
    if object_path:
        guessed = os.path.splitext(str(object_path))[1]
    if not guessed and url:
        guessed = os.path.splitext(str(url).split("?", 1)[0])[1]
    if guessed:
        return guessed
    mime = str(content_type or "").strip()
    if mime:
        ext = mimetypes.guess_extension(mime, strict=False) or ""
        if ext:
            return ext
    return ".png"


def _center_last_picture(doc: Document) -> None:
    try:
        paragraph = doc.paragraphs[-1]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
    except Exception:
        return


def _add_picture_resilient(doc: Document, image_path: str, *, width: Any) -> None:
    try:
        doc.add_picture(image_path, width=width)
        return
    except Exception:
        pass

    normalized_path = ""
    try:
        import tempfile
        from PIL import Image, ImageOps

        with Image.open(image_path) as image:
            image = ImageOps.exif_transpose(image)
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                normalized_path = tmp.name
            image.save(normalized_path, format="PNG")
        doc.add_picture(normalized_path, width=width)
    finally:
        if normalized_path:
            try:
                os.unlink(normalized_path)
            except Exception:
                pass


def _setup_styles(doc: Document):
    """
    Configure professional styles for the document.
    """
    styles = doc.styles

    # Normal Style (Body Text)
    # Target: 12pt (approx 16px), 1.5 spacing, standard font
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Consistent font for all scripts
    element = font._element.rPr.rFonts
    element.set(qn('w:eastAsia'), 'Microsoft YaHei')
    element.set(qn('w:ascii'), 'Microsoft YaHei')
    element.set(qn('w:hAnsi'), 'Microsoft YaHei')
    
    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing = 1.25
    paragraph_format.space_after = Pt(6)
    # Professional justified alignment for both EN and ZH
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1
    style_h1 = styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Microsoft YaHei'
    font_h1.size = Pt(24)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(0, 0, 0)
    element = font_h1._element.rPr.rFonts
    element.set(qn('w:eastAsia'), 'Microsoft YaHei')
    element.set(qn('w:ascii'), 'Microsoft YaHei')
    element.set(qn('w:hAnsi'), 'Microsoft YaHei')
    style_h1.paragraph_format.space_before = Pt(14)
    style_h1.paragraph_format.space_after = Pt(10)

    # Heading 2
    style_h2 = styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Microsoft YaHei'
    font_h2.size = Pt(18)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(0, 0, 0)
    element = font_h2._element.rPr.rFonts
    element.set(qn('w:eastAsia'), 'Microsoft YaHei')
    element.set(qn('w:ascii'), 'Microsoft YaHei')
    element.set(qn('w:hAnsi'), 'Microsoft YaHei')
    style_h2.paragraph_format.space_before = Pt(12)
    style_h2.paragraph_format.space_after = Pt(6)
    
    # Heading 3
    style_h3 = styles['Heading 3']
    font_h3 = style_h3.font
    font_h3.name = 'Microsoft YaHei'
    font_h3.size = Pt(14)
    font_h3.bold = True
    font_h3.color.rgb = RGBColor(0, 0, 0)
    element = font_h3._element.rPr.rFonts
    element.set(qn('w:eastAsia'), 'Microsoft YaHei')
    element.set(qn('w:ascii'), 'Microsoft YaHei')
    element.set(qn('w:hAnsi'), 'Microsoft YaHei')
    style_h3.paragraph_format.space_before = Pt(10)
    style_h3.paragraph_format.space_after = Pt(2)

    # Heading 4
    style_h4 = styles['Heading 4']
    font_h4 = style_h4.font
    font_h4.name = 'Microsoft YaHei'
    font_h4.size = Pt(12)
    font_h4.bold = True
    font_h4.color.rgb = RGBColor(0, 0, 0)
    element = font_h4._element.rPr.rFonts
    element.set(qn('w:eastAsia'), 'Microsoft YaHei')
    element.set(qn('w:ascii'), 'Microsoft YaHei')
    element.set(qn('w:hAnsi'), 'Microsoft YaHei')
    style_h4.paragraph_format.space_before = Pt(6)
    style_h4.paragraph_format.space_after = Pt(0)


def _add_cover_page(doc: Document, title: str, abstract: str):
    """
    Create a clean, centered cover page.
    """
    # Vertical spacing for visually centering content
    for _ in range(8):
        doc.add_paragraph("")
    
    # Title
    p_title = doc.add_paragraph(title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.runs[0]
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    _force_paragraph_font(p_title)
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    
    # Abstract
    if abstract:
        p_abs = doc.add_paragraph(abstract)
        p_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_abs.style = 'Normal'
        # Override italic or lighter color for abstract
        run_abs = p_abs.runs[0]
        run_abs.font.size = Pt(14)
        run_abs.font.italic = True
        run_abs.font.color.rgb = RGBColor(100, 100, 100)
        _force_paragraph_font(p_abs)
    
    doc.add_page_break()


def _add_toc_page(doc: Document, toc_title: str, entries: List[Tuple[int, str, str]]):
    """
    Insert a Table of Contents. 
    1. A physical fallback list with DOTS and PAGEREF fields.
    2. A Word Field Code (for professional auto-updating).
    """
    toc_p = doc.add_heading(toc_title, level=1)
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _force_paragraph_font(toc_p)
    
    # 1. Physical Fallback (List of headers with dots and PAGEREF)
    for level, text, bookmark_name in entries:
        if level > 3: continue 
        indent = (level - 1) * 0.5
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        p.add_run("\t") 
        
        # Add PAGEREF field using separate runs for correct OOXML sequence
        # <begin> <instr> <separate> <text> <end>
        # w:dirty="true" forces Word/WPS to recalculate the page number when
        # the document is opened. Without it the placeholder text "1" is
        # shown verbatim for every entry until the user manually invokes
        # "Update Field" — which is exactly the bug we are fixing.
        run_begin = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        fldChar1.set(qn('w:dirty'), 'true')
        run_begin._element.append(fldChar1)
        
        run_instr = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' PAGEREF {bookmark_name} \\h '
        run_instr._element.append(instrText)
        
        run_sep = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_sep._element.append(fldChar2)
        
        # Placeholder text (Word will replace this when updating fields)
        run_val = p.add_run("1") 
        run_val.font.size = Pt(10.5)
        
        run_end = p.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_end._element.append(fldChar3)
        
        _force_paragraph_font(p)

    doc.add_paragraph("") 

    # 2. Word Field Code
    # Same w:dirty="true" trick — without it, the auto TOC field never
    # populates with real page numbers on first open.
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    fldChar.set(qn('w:dirty'), 'true')
    run._element.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._element.append(instrText)
    fldRef = OxmlElement('w:fldChar')
    fldRef.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldRef)
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run._element.append(fldEnd)


def _add_bookmark(paragraph: Paragraph, name: str, bookmark_id: int):
    """
    Wrap paragraph text in a bookmark.
    """
    p_element = paragraph._p
    
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), name)
    
    # Insert at the beginning of the paragraph element
    p_element.insert(0, bookmark_start)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    p_element.append(bookmark_end)


def _enable_update_fields_on_open(doc: Document):
    """
    Ask Word to update fields (e.g., TOC) when the document is opened.
    """
    settings = doc._part.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _process_paragraph_text(paragraph: Paragraph, text: str):
    """
    Process inline formatting like **bold** in a paragraph using regex.
    """
    # Simple bold parser: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')


def _force_paragraph_font(paragraph: Paragraph):
    for run in paragraph.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')


def generate_docx_file(
    markdown_content: str,
    filename: str = "report.docx",
    template_path: Optional[str] = None,
    return_metadata: bool = False,
    cover_title: Optional[str] = None,
    skip_cover: bool = False,
    skip_toc: bool = False,
) -> str | Dict[str, Any]:
    """
    Generate a professional DOCX report from Markdown.
    Includes Cover Page, Table of Contents, and Styled Body.

    skip_cover: when True, no cover page is rendered. Set by upstream for
        short-form deliverables (resume, email, memo) where a standalone
        cover page is editorial noise.
    skip_toc: when True, the in-document table of contents is suppressed
        even if the markdown explicitly contains a "## 目录" heading.
        Mirrors skip_cover semantics for the same short-form cases.
    """
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
    _setup_styles(doc)

    markdown_content = _normalize_image_markers(markdown_content)
    lines = markdown_content.splitlines()
    language = "zh" if re.search(r"[\u4e00-\u9fff]", markdown_content or "") else "en"
    toc_title = i18n_get_list("toc_heading_variants", language)[0]

    # 1. Parse Structure
    # State 0: Cover, State 1: TOC (ignore), State 2: Body
    has_manual_toc = False

    # 优先使用上游传入的 cover_title（基于 goal_contract/uploaded filename 等解析出的标题）
    resolved_cover_title = str(cover_title or "").strip()
    if not resolved_cover_title:
        resolved_cover_title = i18n_get("cover_title_default", language)
    cover_abstract = ""
    body_lines = []

    # Pre-scan for markdown structure.
    # 若 markdown 根本没有 heading（# / ##），不要走 state machine——所有内容直接作为 body 渲染，
    # 避免把第一段纯文本误判为 abstract 导致正文为空。
    has_any_heading = any(
        line.strip().startswith("# ") or line.strip().startswith("## ") or line.strip().startswith("### ")
        for line in lines
    )
    has_h1_heading = any(line.strip().startswith("# ") for line in lines)
    # 若未显式传入 cover_title，尝试用 markdown 第一个 H1 作为封面标题
    if not cover_title:
        for line in lines:
            if line.strip().startswith("# "):
                resolved_cover_title = line.strip()[2:]
                break

    if not has_any_heading:
        # 没有任何标题：所有内容都进 body，不做 cover abstract 抽取
        state = 2
        for line in lines:
            body_lines.append(line)
    else:
        state = 0

    # Process lines (仅当有 heading 时进入 state machine)
    for line in lines if has_any_heading else []:
        stripped = line.strip()

        # Structure Detection
        if state == 0:
            toc_tokens = i18n_get_list("toc_heading_variants", "zh") + i18n_get_list("toc_heading_variants", "en")
            toc_pattern = r"^#{1,2}\s+(" + "|".join(re.escape(t) for t in toc_tokens) + r")"
            if re.match(toc_pattern, stripped, re.IGNORECASE):
                has_manual_toc = True
                state = 1
                continue
            if stripped.startswith("# "):
                if skip_cover:
                    state = 2
                    body_lines.append(line)
                continue
            # Accumulate abstract content if not title
            if stripped and not stripped.startswith("#"):
                # Only treat leading non-heading text as a cover abstract when
                # the source markdown provides a document-level H1. Otherwise
                # this is normal body content and must not be swallowed into the
                # skipped cover page.
                if has_h1_heading and not skip_cover:
                    if not cover_abstract:
                        cover_abstract = stripped
                else:
                    state = 2
                    body_lines.append(line)
            elif stripped.startswith(("## ", "### ", "#### ")):
                # Jump straight to body if we missed TOC
                state = 2
                body_lines.append(line)
        
        elif state == 1:
            if stripped == "---":
                state = 2
                continue
            if stripped.startswith("## ") or stripped.startswith("# "):
                 # Transition to body if we hit any real header
                 state = 2
                 body_lines.append(line)
            continue # Ignore markdown TOC
            
        elif state == 2:
            # Ignore leading spacers/empty lines at start of body
            if not body_lines and not stripped:
                continue
            if stripped == "---":
                 # Preserve horizontal separators as visual separators, not page breaks.
                 body_lines.append("__HORIZONTAL_RULE__")
                 continue
            body_lines.append(line)

    # 2. Render Cover (gated for short-form deliverables)
    if not skip_cover:
        _add_cover_page(doc, resolved_cover_title, cover_abstract)

    # 3. Pre-scan Body for TOC Entries & Assign Bookmarks
    toc_entries = []
    bookmark_counter = 0
    for line in body_lines:
        s = line.strip()
        if s.startswith("# "):
            level = 1
            text = s[2:].strip()
        elif s.startswith("## "):
            level = 2
            text = s[3:].strip()
        elif s.startswith("### "):
            level = 3
            text = s[4:].strip()
        else:
            continue
        
        bookmark_name = f"_Toc{bookmark_counter}"
        toc_entries.append((level, text, bookmark_name))
        bookmark_counter += 1

    # 4. Render TOC only when the source markdown explicitly includes one.
    # skip_toc forcibly disables the TOC even if a "## 目录" heading slipped
    # through (e.g. user-edited markdown of a short-form deliverable).
    if has_manual_toc and toc_entries and not skip_toc:
        _add_toc_page(doc, toc_title, toc_entries)
        doc.add_page_break()

    # 5. Render Body
    bookmark_idx = 0
    bullet_re = re.compile(r"^-\s+")
    number_re = re.compile(r"^\d+\.\s+")
    table_rows: List[str] = []
    in_table = False
    
    image_line_count = 0
    image_fetch_ok = 0
    image_fetch_fail = 0
    in_code_block = False
    code_block_type = ""
    pending_empty = False
    for raw_line in body_lines:
        line = raw_line.strip()
        if not line:
            # Skip empty lines to reduce excessive spacing in the document.
            # Markdown empty lines between blocks usually don't need dedicated 
            # paragraphs when styles already have space_before/after.
            continue
        pending_empty = False
        if line.startswith("```"):
            if in_code_block:
                in_code_block = False
                code_block_type = ""
                continue
            in_code_block = True
            code_block_type = line.lstrip("`").strip().lower()
            # Insert a placeholder for non-text blocks
            if "chart" in code_block_type or "mermaid" in code_block_type or "graph" in code_block_type:
                p = doc.add_paragraph("[Chart/Diagram]", style="Normal")
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            continue
        if in_code_block:
            # Allow tables inside code blocks to fall through and be rendered
            if '|' in line and (('-' in line and '--' in line) or any(c.strip() for c in line.split('|'))):
                pass
            else:
                continue
        if line == "__PAGE_BREAK__":
            doc.add_page_break()
            continue
        if line == "__HORIZONTAL_RULE__":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_" * 28)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
            continue
        if '|' in line and not in_table and not line.startswith('#'):
            in_table = True
            table_rows.append(line)
            continue
        elif in_table and '|' in line:
            table_rows.append(line)
            continue
        elif in_table:
            # flush table
            rows = [r for r in table_rows if r.strip()]
            if rows:
                # detect header separator
                has_sep = False
                if len(rows) > 1 and re.match(r"^\|?[:\-\s|]+\|?$", rows[1]) and '-' in rows[1]:
                    has_sep = True
                start_idx = 0
                headers = []
                if has_sep:
                    headers = [c.strip() for c in rows[0].strip().strip('|').split('|')]
                    start_idx = 2
                data_rows = rows[start_idx:]
                cols = len(headers) if headers else max((len(r.strip().strip('|').split('|')) for r in data_rows), default=1)
                table = doc.add_table(rows=1 if headers else 0, cols=cols)
                table.style = "Table Grid"
                if headers:
                    hdr_cells = table.rows[0].cells
                    for i, h in enumerate(headers):
                        if i < len(hdr_cells):
                            hdr_cells[i].text = h
                for r in data_rows:
                    cells = [c.strip() for c in r.strip().strip('|').split('|')]
                    row_cells = table.add_row().cells
                    for i, c in enumerate(cells):
                        if i < len(row_cells):
                            row_cells[i].text = c
            table_rows = []
            in_table = False
            
        if line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            _force_paragraph_font(p)
        elif line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            _add_bookmark(p, f"_Toc{bookmark_idx}", bookmark_idx)
            bookmark_idx += 1
            _force_paragraph_font(p)
        elif line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            _add_bookmark(p, f"_Toc{bookmark_idx}", bookmark_idx)
            bookmark_idx += 1
            _force_paragraph_font(p)
        elif line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            _add_bookmark(p, f"_Toc{bookmark_idx}", bookmark_idx)
            bookmark_idx += 1
            _force_paragraph_font(p)
        elif bullet_re.match(line):
            text = bullet_re.sub("", line)
            p = doc.add_paragraph(style="List Bullet")
            _process_paragraph_text(p, text)
            _force_paragraph_font(p)
        elif number_re.match(line):
            # Render numbered lists as plain paragraphs to avoid Word's automatic 
            # list levels being miscalculated or restarted.
            p = doc.add_paragraph(style="Normal")
            _process_paragraph_text(p, line)
            _force_paragraph_font(p)
        elif line.startswith("![") or "<img" in line.lower():
            image_line_count += 1
            match = IMAGE_RE.match(line)
            img_url = ""
            if match:
                img_url = _extract_image_url(match.group(1))
            else:
                html_match = HTML_IMG_RE.search(line)
                if html_match:
                    img_url = _extract_image_url(html_match.group(1))
                elif "(" in line and ")" in line:
                    # Fallback extraction if markdown regex fails
                    img_url = _extract_image_url(line.split("(", 1)[1].rsplit(")", 1)[0])

            if not img_url:
                image_fetch_fail += 1
                p = doc.add_paragraph("[Image/Chart Component]", style="Normal")
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
                continue
            img_url = _resolve_render_image_url(img_url)
            try:
                import base64
                import binascii
                import os as _os
                import tempfile

                if img_url.startswith("data:image/"):
                    try:
                        header, b64data = img_url.split(",", 1)
                        ext = header.split(";")[0].split("/")[1]
                        img_bytes = base64.b64decode(b64data)
                        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
                            tmp.write(img_bytes)
                            tmp_path = tmp.name
                        _add_picture_resilient(doc, tmp_path, width=Cm(14))
                        _center_last_picture(doc)
                        os.unlink(tmp_path)
                        image_fetch_ok += 1
                        continue
                    except (ValueError, binascii.Error):
                        raise RuntimeError("invalid data image")

                if img_url.startswith("file://"):
                    img_url = img_url[7:]
                if _os.path.exists(img_url):
                    _add_picture_resilient(doc, img_url, width=Cm(14))
                    _center_last_picture(doc)
                    image_fetch_ok += 1
                    continue

                uploader = AliyunOSSUploader()
                object_path = uploader.object_path_from_url(img_url)
                if object_path:
                    local_path = ""
                    try:
                        if uploader.storage_backend == "local":
                            local_path = str(uploader.local_file_path(object_path))
                    except Exception:
                        local_path = ""
                    if local_path and _os.path.exists(local_path):
                        _add_picture_resilient(doc, local_path, width=Cm(14))
                        _center_last_picture(doc)
                        image_fetch_ok += 1
                        continue
                    img_bytes = uploader.read_bytes(object_path)
                    suffix = _guess_image_suffix(
                        url=img_url,
                        object_path=object_path,
                        content_type=uploader.guess_content_type(object_path),
                    )
                    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix or ".png") as tmp:
                        tmp.write(img_bytes)
                        tmp_path = tmp.name
                    _add_picture_resilient(doc, tmp_path, width=Cm(14))
                    _center_last_picture(doc)
                    os.unlink(tmp_path)
                    image_fetch_ok += 1
                    continue

                import httpx
                headers = {
                    "User-Agent": "Mozilla/5.0",
                    "Accept": "image/*,*/*;q=0.8",
                }
                with httpx.Client(timeout=20.0, follow_redirects=True) as client:
                    resp = client.get(img_url, headers=headers)
                if resp.status_code == 200:
                    suffix = _guess_image_suffix(
                        url=img_url,
                        content_type=str(resp.headers.get("content-type") or "").strip(),
                    )
                    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix or ".png") as tmp:
                        tmp.write(resp.content)
                        tmp_path = tmp.name
                    _add_picture_resilient(doc, tmp_path, width=Cm(14))
                    _center_last_picture(doc)
                    os.unlink(tmp_path)
                    image_fetch_ok += 1
                else:
                    raise RuntimeError("image fetch failed")
            except Exception as exc:
                log_print(f"[docx-image] failed type={type(exc).__name__} url={img_url[:160]}")
                image_fetch_fail += 1
                p = doc.add_paragraph("[Image/Chart Component]", style="Normal")
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        else:
            p = doc.add_paragraph(style="Normal")
            # Apply first-line indent if the text is long enough to be a real paragraph
            if len(line.strip()) > 10 and language == "zh":
                # Standard indent for Chinese is ~2 characters (approx 0.7cm - 1cm depending on font)
                p.paragraph_format.first_line_indent = Pt(24) 
            _process_paragraph_text(p, line)
            _force_paragraph_font(p)

    # Flush trailing table if markdown ends with table
    if in_table and table_rows:
        rows = [r for r in table_rows if r.strip()]
        if rows:
            has_sep = False
            if len(rows) > 1 and re.match(r"^\|?[:\-\s|]+\|?$", rows[1]) and '-' in rows[1]:
                has_sep = True
            start_idx = 0
            headers = []
            if has_sep:
                headers = [c.strip() for c in rows[0].strip().strip('|').split('|')]
                start_idx = 2
            data_rows = rows[start_idx:]
            cols = len(headers) if headers else max((len(r.strip().strip('|').split('|')) for r in data_rows), default=1)
            table = doc.add_table(rows=1 if headers else 0, cols=cols)
            table.style = "Table Grid"
            if headers:
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = h
            for r in data_rows:
                cells = [c.strip() for c in r.strip().strip('|').split('|')]
                row_cells = table.add_row().cells
                for i, c in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = c

    log_print(f"[docx-image] lines={image_line_count} ok={image_fetch_ok} fail={image_fetch_fail}")

    # Ensure Word updates TOC fields on open
    _enable_update_fields_on_open(doc)

    # Save
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    static_dir = os.path.join(base_dir, "static")
    os.makedirs(static_dir, exist_ok=True)

    if not filename.endswith(".docx"):
        filename = f"{filename}.docx"
    file_path = os.path.join(static_dir, filename)
    doc.save(file_path)
    if not return_metadata:
        return file_path
    return {
        "file_path": file_path,
        "render_stats": {
            "image_line_count": image_line_count,
            "image_fetch_ok": image_fetch_ok,
            "image_fetch_fail": image_fetch_fail,
        },
    }


@tool
def create_docx_from_markdown(markdown_content: str, filename: str = "report.docx") -> str:
    """
    Generate a Word document (.docx) from markdown-like text.
    """
    try:
        file_path = generate_docx_file(markdown_content, filename)
        return f"DOCX Generated Successfully: {file_path}"
    except Exception as exc:
        return f"DOCX Generation Failed: {str(exc)}"
