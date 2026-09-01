from __future__ import annotations

import json
import logging
import re
import sys
import tempfile
from importlib import util as importlib_util
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from app.domain.jobs import ParseDocumentJobRequest
from app.integrations.callbacks.admin_api_client import post_callback
from app.integrations.converters.libreoffice import convert_legacy_office_to_modern
from app.integrations.storage import get_storage_adapter
from app.repositories.job_repository import mark_job_succeeded, update_job_progress
from app.services.rag_chunk_optimizer import RagChunkOptions, build_rag_chunks

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = {"txt", "md", "markdown", "csv", "json"}
DOCX_EXTENSIONS = {"docx"}
PDF_EXTENSIONS = {"pdf"}
LEGACY_OFFICE_EXTENSIONS = {"doc", "ppt", "xls"}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}


@dataclass
class ParsedDocument:
    markdown: str
    raw: dict[str, Any]
    raw_chunks: list[dict[str, Any]] | None = None
    rag_chunks: list[dict[str, Any]] | None = None


def run_document_parse(job_id: str, request: ParseDocumentJobRequest) -> None:
    source_storage = get_storage_adapter(request.source.storageType)
    artifact_storage = get_storage_adapter(request.artifacts.storageType)
    artifact_prefix = request.artifacts.storagePrefix.strip().strip("/")

    with tempfile.TemporaryDirectory(prefix="askai-parse-") as temp_dir:
        temp_path = Path(temp_dir)
        suffix = Path(request.source.filename or request.source.storageKey).suffix or ".document"
        source_path = temp_path / f"source{suffix}"
        with source_storage.open_file(request.source.storageKey) as source, source_path.open("wb") as output:
            while True:
                chunk = source.read(1024 * 1024)
                if not chunk:
                    break
                output.write(chunk)

        update_job_progress(job_id, 25)
        parsed = parse_document(source_path, request.source.filename or request.source.storageKey)
        update_job_progress(job_id, 55)
        raw_chunks = parsed.raw_chunks if parsed.raw_chunks is not None else chunk_markdown(parsed.markdown, request.chunkSize, request.chunkOverlap)
        rag_chunks = parsed.rag_chunks if parsed.rag_chunks is not None else build_rag_chunks(
            raw_chunks,
            _rag_chunk_options(request.minChunkSize, request.chunkSize, request.chunkOverlap),
        )
        update_job_progress(job_id, 75)

        markdown_key = f"{artifact_prefix}/parsed.md"
        json_key = f"{artifact_prefix}/parsed.json"
        raw_chunks_key = f"{artifact_prefix}/raw_chunks.json"
        rag_chunks_key = f"{artifact_prefix}/rag_chunks.json"
        chunks_key = f"{artifact_prefix}/chunks.json"

        _write_text_artifact(artifact_storage, markdown_key, parsed.markdown)
        _write_json_artifact(artifact_storage, json_key, parsed.raw)
        _write_json_artifact(artifact_storage, raw_chunks_key, {"chunkStage": "raw", "chunks": raw_chunks})
        _write_json_artifact(artifact_storage, rag_chunks_key, {"chunkStage": "rag", "chunks": rag_chunks})
        _write_json_artifact(artifact_storage, chunks_key, {"chunkStage": "rag", "chunks": rag_chunks})

    result = {
        "markdownKey": markdown_key,
        "jsonKey": json_key,
        "rawChunksKey": raw_chunks_key,
        "rawChunkCount": len(raw_chunks),
        "ragChunksKey": rag_chunks_key,
        "ragChunkCount": len(rag_chunks),
        "chunksKey": chunks_key,
        "chunkCount": len(rag_chunks),
    }
    update_job_progress(job_id, 90)
    post_callback(
        request.callback.url,
        request.callback.token,
        {
            "jobId": job_id,
            "status": "succeeded",
            **result,
            "error": "",
        },
    )
    mark_job_succeeded(job_id, result)


def parse_document(path: Path, filename: str) -> ParsedDocument:
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext in LEGACY_OFFICE_EXTENSIONS:
        converted_path = convert_legacy_office_to_modern(path, path.parent)
        path = converted_path
        ext = converted_path.suffix.lower().lstrip(".")
        filename = f"{Path(filename).stem}{converted_path.suffix}"
    if ext in IMAGE_EXTENSIONS:
        return parse_image_with_ocr(path, ext)

    try:
        from docling.document_converter import DocumentConverter  # type: ignore
    except ImportError:
        return parse_with_fallback(path, ext)
    else:
        converter = _build_docling_converter(ext)
        try:
            result = converter.convert(str(path))
        except Exception:
            if ext in TEXT_EXTENSIONS:
                return parse_plain_text(path, ext)
            raise
        document = result.document
        markdown = document.export_to_markdown()
        try:
            raw = document.export_to_dict()
        except Exception:
            raw = {"parser": "docling", "markdown": markdown}
        image_ocr_pages = _pdf_image_ocr_pages(path) if ext == "pdf" else set()
        raw_chunks = chunk_docling_document(document, image_ocr_pages=image_ocr_pages)
        return ParsedDocument(markdown=markdown, raw=raw, raw_chunks=raw_chunks)


def _build_docling_converter(ext: str) -> Any:
    if ext != "pdf":
        from docling.document_converter import DocumentConverter  # type: ignore

        return DocumentConverter()

    if importlib_util.find_spec("onnxruntime") is None:
        from docling.document_converter import DocumentConverter  # type: ignore

        logger.warning(
            "onnxruntime is not installed in %s; using Docling default OCR for PDF parsing",
            sys.executable,
        )
        return DocumentConverter()

    try:
        from docling.datamodel.base_models import InputFormat  # type: ignore
        from docling.datamodel.pipeline_options import RapidOcrOptions  # type: ignore
        from docling.document_converter import DocumentConverter, PdfFormatOption  # type: ignore
        from app.services.docling_runtime import pdf_pipeline_options
    except Exception:
        from docling.document_converter import DocumentConverter  # type: ignore

        return DocumentConverter()

    pipeline_options = pdf_pipeline_options()
    pipeline_options.do_ocr = True
    pipeline_options.ocr_options = RapidOcrOptions(
        force_full_page_ocr=True,
        lang=["chinese", "english"],
        backend="onnxruntime",
    )
    logger.info("Using RapidOCR/onnxruntime for PDF parsing")
    return DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
        }
    )


def chunk_docling_document(document: Any, *, image_ocr_pages: set[int] | None = None) -> list[dict[str, Any]]:
    try:
        from docling.chunking import HybridChunker  # type: ignore
        from docling_core.transforms.chunker.tokenizer.base import BaseTokenizer  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Docling chunker 不可用，无法使用 Docling 分段") from exc

    class AskAILocalTokenizer(BaseTokenizer):  # type: ignore[misc]
        max_tokens: int = 900

        def count_tokens(self, text: str) -> int:
            return _estimate_token_count(text)

        def get_max_tokens(self) -> int:
            return self.max_tokens

        def get_tokenizer(self) -> Any:
            return self.count_tokens

    chunker = HybridChunker(
        tokenizer=AskAILocalTokenizer(),
        merge_peers=True,
        repeat_table_header=True,
    )
    chunks: list[dict[str, Any]] = []
    emitted_table_refs: set[str] = set()
    for index, chunk in enumerate(chunker.chunk(document)):
        text = _clean_chunk_text(str(getattr(chunk, "text", "") or ""))
        if not text or _is_noise_chunk(text):
            continue
        chunk_meta = getattr(chunk, "meta", None)
        meta = _export_docling_model(chunk_meta)
        raw_chunk = _export_docling_model(chunk)
        title_path = _docling_title_path(meta)
        page_no = _docling_page_no(meta)
        content_type = _docling_content_type(meta)
        source_kind = "image_ocr" if page_no in (image_ocr_pages or set()) else ""
        source_anchor = _docling_source_anchor(chunk_meta, page_no) or _docling_source_anchor(meta, page_no)
        table_chunks = _docling_table_row_chunks(
            document=document,
            chunk=chunk,
            meta=meta,
            raw_chunk=raw_chunk,
            title_path=title_path,
            page_no=page_no,
            start_ordinal=len(chunks),
            emitted_table_refs=emitted_table_refs,
        )
        if table_chunks:
            chunks.extend(table_chunks)
            if content_type == "table" or _looks_like_flattened_table_text(text):
                continue
        try:
            contextual_text = _clean_chunk_text(str(chunker.contextualize(chunk) or text))
        except Exception:
            contextual_text = "\n".join(title_path + [text]) if title_path else text
        chunks.append(
            {
                "chunkId": f"raw_{len(chunks) + 1:06d}",
                "ordinal": len(chunks),
                "text": text,
                "contextualText": contextual_text,
                "titlePath": title_path,
                "pageNo": page_no,
                "contentType": content_type,
                "chunkStage": "raw",
                "sourceChunkIds": [],
                "metadata": {
                    "parser": "docling",
                    **({"sourceKind": source_kind} if source_kind else {}),
                    "chunker": "HybridChunker",
                    "tokenizer": "AskAILocalTokenizer",
                    "doclingMeta": meta,
                    "doclingChunk": raw_chunk,
                    "sourceAnchor": source_anchor,
                },
            }
        )
    return chunks


def _pdf_image_ocr_pages(path: Path) -> set[int]:
    """Identify scanned pages that have no meaningful native PDF text layer."""
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
    except Exception:
        return set()

    image_pages: set[int] = set()
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            native_text = re.sub(r"\s+", "", page.extract_text() or "")
        except Exception:
            native_text = ""
        if len(native_text) < 20:
            image_pages.add(page_number)
    return image_pages


def _docling_table_row_chunks(
    *,
    document: Any,
    chunk: Any,
    meta: dict[str, Any],
    raw_chunk: dict[str, Any],
    title_path: list[str],
    page_no: int | None,
    start_ordinal: int,
    emitted_table_refs: set[str],
) -> list[dict[str, Any]]:
    table_items = _docling_table_items(document, chunk, meta)
    output: list[dict[str, Any]] = []
    for table_item in table_items:
        table_ref = str(getattr(table_item, "self_ref", "") or "")
        if table_ref and table_ref in emitted_table_refs:
            continue
        try:
            row_chunks = _table_item_to_row_chunks(
                document=document,
                table_item=table_item,
                meta=meta,
                raw_chunk=raw_chunk,
                title_path=title_path,
                page_no=page_no,
                start_ordinal=start_ordinal + len(output),
            )
        except Exception:
            continue
        if row_chunks:
            if table_ref:
                emitted_table_refs.add(table_ref)
            output.extend(row_chunks)
    return output


def _docling_table_items(document: Any, chunk: Any, exported_meta: dict[str, Any]) -> list[Any]:
    seen_refs: set[str] = set()
    output: list[Any] = []

    meta = getattr(chunk, "meta", None)
    for item in list(getattr(meta, "doc_items", None) or []):
        if not _is_docling_table_item(item):
            continue
        table_ref = str(getattr(item, "self_ref", "") or "")
        if table_ref and table_ref in seen_refs:
            continue
        if table_ref:
            seen_refs.add(table_ref)
        output.append(item)

    for item in exported_meta.get("doc_items") or []:
        if not isinstance(item, dict):
            continue
        if _docling_label_text(item.get("label") or "") != "table":
            continue
        table_ref = str(item.get("self_ref") or "").strip()
        if not table_ref or table_ref in seen_refs:
            continue
        table_item = _resolve_docling_ref(document, table_ref)
        if _is_docling_table_item(table_item):
            seen_refs.add(table_ref)
            output.append(table_item)
    return output


def _is_docling_table_item(item: Any) -> bool:
    label = _docling_label_text(getattr(item, "label", "")).lower()
    return label == "table" and hasattr(getattr(item, "data", None), "grid")


def _resolve_docling_ref(document: Any, ref: str) -> Any:
    parts = str(ref or "").split("/")
    if len(parts) == 3 and parts[0] == "#":
        collection_name = parts[1]
        try:
            index = int(parts[2])
        except (TypeError, ValueError):
            return None
        collection = getattr(document, collection_name, None)
        if isinstance(collection, list) and 0 <= index < len(collection):
            return collection[index]
    return None


def _table_item_to_row_chunks(
    *,
    document: Any,
    table_item: Any,
    meta: dict[str, Any],
    raw_chunk: dict[str, Any],
    title_path: list[str],
    page_no: int | None,
    start_ordinal: int,
) -> list[dict[str, Any]]:
    table_data = getattr(table_item, "data", None)
    grid = list(getattr(table_data, "grid", None) or [])
    if not grid:
        return []

    rows = [[_table_cell_text(cell, document) for cell in row] for row in grid]
    rows = [[cell.strip() for cell in row] for row in rows]
    if not rows or not any(any(cell for cell in row) for row in rows):
        return []

    header_count = _table_header_row_count(grid)
    if header_count <= 0:
        header_count = 1
    headers = _headers_from_rows(rows[:header_count])
    body_rows = rows[header_count:] if len(rows) > header_count else rows
    if not body_rows:
        return []

    table_ref = str(getattr(table_item, "self_ref", "") or "")
    num_rows = int(getattr(table_data, "num_rows", 0) or len(rows))
    num_cols = int(getattr(table_data, "num_cols", 0) or len(headers))
    output: list[dict[str, Any]] = []
    for row_offset, row in enumerate(body_rows, start=header_count):
        if not any(cell.strip() for cell in row):
            continue
        markdown = _single_row_markdown(headers, row)
        if not markdown:
            continue
        ordinal = start_ordinal + len(output)
        table_context = {
            "tableRef": table_ref,
            "rowIndex": row_offset,
            "rowOrdinal": len(output),
            "headerRows": header_count,
            "numRows": num_rows,
            "numCols": num_cols,
            "displayMode": "group_by_table_ref",
        }
        grid_row = grid[row_offset] if 0 <= row_offset < len(grid) else []
        source_anchor = _table_row_source_anchor(grid_row, page_no) if grid_row else None
        if source_anchor is None:
            source_anchor = _docling_source_anchor(table_item, page_no) or _docling_source_anchor(meta, page_no)
        if source_anchor:
            source_anchor["tableContext"] = table_context
        output.append(
            {
                "chunkId": f"raw_{ordinal + 1:06d}",
                "ordinal": ordinal,
                "text": markdown,
                "contextualText": "\n".join(title_path + [markdown]) if title_path else markdown,
                "titlePath": title_path,
                "pageNo": page_no,
                "contentType": "table_row",
                "chunkStage": "raw",
                "sourceChunkIds": [],
                "metadata": {
                    "parser": "docling",
                    "chunker": "HybridChunker",
                    "tokenizer": "AskAILocalTokenizer",
                    "doclingMeta": meta,
                    "doclingChunk": raw_chunk,
                    "tableContext": table_context,
                    "sourceAnchor": source_anchor,
                },
            }
        )
    return output


def _table_header_row_count(grid: list[list[Any]]) -> int:
    count = 0
    for row in grid:
        if any(bool(getattr(cell, "column_header", False)) for cell in row):
            count += 1
            continue
        break
    return count


def _headers_from_rows(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    width = max((len(row) for row in rows), default=0)
    headers: list[str] = []
    for col_index in range(width):
        parts = []
        for row in rows:
            value = row[col_index].strip() if col_index < len(row) else ""
            if value and value not in parts:
                parts.append(value)
        header = " / ".join(parts).strip() or f"Column {col_index + 1}"
        headers.append(header)
    return headers


def _single_row_markdown(headers: list[str], row: list[str]) -> str:
    if not headers:
        return ""
    width = max(len(headers), len(row))
    safe_headers = [_markdown_table_cell(headers[index] if index < len(headers) else f"Column {index + 1}") for index in range(width)]
    safe_row = [_markdown_table_cell(row[index] if index < len(row) else "") for index in range(width)]
    separator = ["---"] * width
    return "\n".join(
        [
            "| " + " | ".join(safe_headers) + " |",
            "| " + " | ".join(separator) + " |",
            "| " + " | ".join(safe_row) + " |",
        ]
    )


def _markdown_table_cell(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text.replace("|", "\\|")


def _looks_like_flattened_table_text(text: str) -> bool:
    value = str(text or "")
    assignments = re.findall(r"(?:^|[,;])\s*[^,;=]{1,80}\s*=", value)
    return len(assignments) >= 2


def _table_cell_text(cell: Any, document: Any) -> str:
    try:
        if hasattr(cell, "_get_text"):
            return str(cell._get_text(doc=document) or "")
    except Exception:
        pass
    return str(getattr(cell, "text", "") or "")


def _docling_source_anchor(meta: Any, fallback_page_no: int | None = None) -> dict[str, Any]:
    bboxes: list[dict[str, Any]] = []
    page_no = fallback_page_no
    for item in _docling_anchor_items(meta):
        for prov in _docling_item_provs(item):
            prov_page_no = _get_docling_field(prov, "page_no")
            if page_no is None and prov_page_no is not None:
                try:
                    page_no = int(prov_page_no)
                except (TypeError, ValueError):
                    page_no = None
            bbox = _normalize_docling_bbox(_get_docling_field(prov, "bbox"))
            if bbox:
                bboxes.append(bbox)
    return _make_source_anchor(page_no, bboxes)


def _docling_anchor_items(meta: Any) -> list[Any]:
    if meta is None:
        return []
    if isinstance(meta, dict):
        direct_items = meta.get("doc_items")
        if isinstance(direct_items, list):
            return direct_items
        if meta.get("prov") or meta.get("bbox"):
            return [meta]
        return []
    direct_items = getattr(meta, "doc_items", None)
    if isinstance(direct_items, list):
        return direct_items
    if getattr(meta, "prov", None) or getattr(meta, "bbox", None):
        return [meta]
    return []


def _docling_item_provs(item: Any) -> list[Any]:
    if isinstance(item, dict):
        provs = item.get("prov")
        if isinstance(provs, list) and provs:
            return provs
        if item.get("bbox"):
            return [{"bbox": item.get("bbox"), "page_no": item.get("page_no")}]
        return []
    provs = getattr(item, "prov", None)
    if isinstance(provs, list) and provs:
        return provs
    bbox = getattr(item, "bbox", None)
    if bbox:
        return [{"bbox": bbox, "page_no": getattr(item, "page_no", None)}]
    return []


def _get_docling_field(value: Any, field: str) -> Any:
    if isinstance(value, dict):
        return value.get(field)
    return getattr(value, field, None)


def _table_row_source_anchor(row: list[Any], fallback_page_no: int | None = None) -> dict[str, Any] | None:
    bboxes = []
    for cell in row:
        bbox = _normalize_docling_bbox(getattr(cell, "bbox", None))
        if bbox:
            bboxes.append(bbox)
    return _make_source_anchor(fallback_page_no, bboxes) or None


def _normalize_docling_bbox(raw: Any) -> dict[str, Any] | None:
    if raw is None:
        return None
    if hasattr(raw, "model_dump"):
        try:
            raw = raw.model_dump(mode="json", by_alias=True, exclude_none=True)
        except Exception:
            return None
    if not isinstance(raw, dict):
        return None
    try:
        left = float(raw.get("l"))
        top = float(raw.get("t"))
        right = float(raw.get("r"))
        bottom = float(raw.get("b"))
    except (TypeError, ValueError):
        return None
    if right <= left or bottom == top:
        return None
    coord_origin = str(raw.get("coord_origin") or raw.get("coordOrigin") or "TOPLEFT").upper()
    return {
        "l": left,
        "t": top,
        "r": right,
        "b": bottom,
        "coordOrigin": coord_origin,
    }


def _make_source_anchor(page_no: int | None, bboxes: list[dict[str, Any]]) -> dict[str, Any]:
    deduped = []
    seen = set()
    for bbox in bboxes:
        key = (
            round(float(bbox["l"]), 3),
            round(float(bbox["t"]), 3),
            round(float(bbox["r"]), 3),
            round(float(bbox["b"]), 3),
            str(bbox.get("coordOrigin") or ""),
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(bbox)
    anchor: dict[str, Any] = {}
    if page_no is not None:
        anchor["pageNo"] = page_no
    if deduped:
        anchor["bboxes"] = deduped
        if len(deduped) == 1:
            anchor["bbox"] = deduped[0]
    return anchor


def _rag_chunk_options(min_chunk_size: int, chunk_size: int, chunk_overlap: int) -> RagChunkOptions:
    max_tokens = max(400, int(chunk_size or 1000))
    min_tokens = max(50, min(int(min_chunk_size or 400), max_tokens))
    return RagChunkOptions(
        min_tokens=min_tokens,
        target_tokens=max(min_tokens, min(max_tokens, int(max_tokens * 0.75))),
        max_tokens=max_tokens,
        overlap_tokens=max(0, min(int(chunk_overlap or 0), max_tokens // 3)),
    )


def _estimate_token_count(text: str) -> int:
    # Lightweight local tokenizer for Docling HybridChunker. CJK chars are close to
    # one token each; ASCII words are grouped so English text is not over-counted.
    if not text:
        return 0
    cjk_chars = len(re.findall(r"[\u3400-\u9fff]", text))
    words = len(re.findall(r"[A-Za-z0-9_]+", text))
    other_chars = len(re.sub(r"[\u3400-\u9fffA-Za-z0-9_\s]", "", text))
    return max(1, cjk_chars + words + max(1, other_chars // 2))


def _clean_chunk_text(text: str) -> str:
    lines = []
    for line in str(text or "").splitlines():
        if _is_noise_chunk(line):
            continue
        lines.append(line.rstrip())
    return "\n".join(lines).strip()


def _is_noise_chunk(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if not compact:
        return True
    if len(compact) <= 3:
        return True
    return bool(re.fullmatch(r"[_\-—=~·.。…、|/\\]+", compact))


def _export_docling_model(value: Any) -> dict[str, Any]:
    if value is None:
        return {}
    try:
        if hasattr(value, "export_json_dict"):
            exported = value.export_json_dict()
        elif hasattr(value, "model_dump"):
            exported = value.model_dump(mode="json", by_alias=True, exclude_none=True)
        else:
            exported = value
    except Exception:
        return {"repr": repr(value)}
    if isinstance(exported, dict):
        return exported
    return {"value": exported}


def _docling_title_path(meta: dict[str, Any]) -> list[str]:
    headings = meta.get("headings")
    if not isinstance(headings, list):
        return []
    return [str(item).strip() for item in headings if str(item).strip()]


def _docling_page_no(meta: dict[str, Any]) -> int | None:
    for item in meta.get("doc_items") or []:
        if not isinstance(item, dict):
            continue
        for prov in item.get("prov") or []:
            if isinstance(prov, dict) and prov.get("page_no") is not None:
                try:
                    return int(prov.get("page_no"))
                except (TypeError, ValueError):
                    return None
    return None


def _docling_content_type(meta: dict[str, Any]) -> str:
    for item in meta.get("doc_items") or []:
        if not isinstance(item, dict):
            continue
        label = item.get("label") or item.get("self_ref") or ""
        if label:
            return _docling_label_text(label)
    return "text"


def _docling_label_text(value: Any) -> str:
    text = str(getattr(value, "value", value) or "").strip()
    if "." in text:
        text = text.rsplit(".", 1)[-1]
    return text.lower()


def parse_with_fallback(path: Path, ext: str) -> ParsedDocument:
    if ext in TEXT_EXTENSIONS:
        return parse_plain_text(path, ext)

    if ext in DOCX_EXTENSIONS:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise RuntimeError("Docling 未安装，且 python-docx 不可用，无法解析 DOCX") from exc
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        markdown = "\n\n".join(paragraphs)
        return ParsedDocument(markdown=markdown, raw={"parser": "python-docx", "paragraphs": paragraphs})

    if ext in PDF_EXTENSIONS:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:
            raise RuntimeError("Docling 未安装，且 pypdf 不可用，无法解析 PDF") from exc
        reader = PdfReader(str(path))
        pages: list[dict[str, Any]] = []
        for index, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page": index + 1, "text": text})
        markdown = "\n\n".join(f"## Page {item['page']}\n\n{item['text']}" for item in pages if item["text"].strip())
        return ParsedDocument(markdown=markdown, raw={"parser": "pypdf", "pages": pages})

    raise RuntimeError("Docling 未安装，当前文件类型没有可用的轻量解析器")


def parse_plain_text(path: Path, ext: str) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(markdown=text, raw={"parser": "text", "format": ext, "text": text})


def parse_image_with_ocr(path: Path, ext: str) -> ParsedDocument:
    if importlib_util.find_spec("rapidocr") is None:
        return ParsedDocument(
            markdown="",
            raw={"parser": "rapidocr", "format": ext, "error": "rapidocr is not installed"},
            raw_chunks=[],
        )
    try:
        from rapidocr import RapidOCR  # type: ignore
    except Exception as exc:
        return ParsedDocument(
            markdown="",
            raw={"parser": "rapidocr", "format": ext, "error": str(exc)},
            raw_chunks=[],
        )

    result = RapidOCR()(path)
    texts = [str(item).strip() for item in getattr(result, "txts", ()) or () if str(item).strip()]
    raw_scores = getattr(result, "scores", None)
    raw_boxes = getattr(result, "boxes", None)
    scores = list(raw_scores) if raw_scores is not None else []
    boxes = list(raw_boxes) if raw_boxes is not None else []
    lines: list[dict[str, Any]] = []
    for index, text in enumerate(texts):
        score = scores[index] if index < len(scores) else None
        box = boxes[index].tolist() if index < len(boxes) and hasattr(boxes[index], "tolist") else (boxes[index] if index < len(boxes) else None)
        lines.append({"text": text, "score": score, "box": box})
    markdown = "\n".join(item["text"] for item in lines).strip()
    raw_chunks = []
    if markdown:
        raw_chunks.append(
            {
                "chunkId": "raw_000001",
                "ordinal": 0,
                "text": markdown,
                "contextualText": markdown,
                "titlePath": [],
                "pageNo": 1,
                "contentType": "text",
                "chunkStage": "raw",
                "sourceChunkIds": [],
                "metadata": {
                    "parser": "rapidocr",
                    "sourceKind": "image_ocr",
                    "format": ext,
                    "lineCount": len(lines),
                    "lines": lines,
                },
            }
        )
    return ParsedDocument(
        markdown=markdown,
        raw={
            "parser": "rapidocr",
            "sourceKind": "image_ocr",
            "format": ext,
            "lineCount": len(lines),
            "lines": lines,
        },
        raw_chunks=raw_chunks,
    )


def chunk_markdown(markdown: str, chunk_size: int, chunk_overlap: int) -> list[dict[str, Any]]:
    blocks = _markdown_blocks(markdown)
    chunks: list[dict[str, Any]] = []
    current_parts: list[str] = []
    current_title_path: list[str] = []
    current_len = 0

    for block in blocks:
        text = block["text"].strip()
        if not text:
            continue
        if block["title_path"]:
            current_title_path = block["title_path"]
        if block.get("type") == "code":
            if current_parts:
                chunks.append(_make_chunk(len(chunks), current_parts, current_title_path))
                current_parts = []
                current_len = 0
            chunks.append(_make_chunk(len(chunks), [text], current_title_path, content_type="code"))
            continue
        if current_parts and current_len + len(text) > chunk_size:
            chunks.append(_make_chunk(len(chunks), current_parts, current_title_path))
            overlap_text = _tail_text_by_lines("\n\n".join(current_parts), chunk_overlap)
            current_parts = [overlap_text] if overlap_text else []
            current_len = len(overlap_text)
        current_parts.append(text)
        current_len += len(text)

    if current_parts:
        chunks.append(_make_chunk(len(chunks), current_parts, current_title_path))
    return chunks


def _markdown_blocks(markdown: str) -> list[dict[str, Any]]:
    title_stack: list[str] = []
    blocks: list[dict[str, Any]] = []
    paragraph: list[str] = []
    code_block: list[str] = []
    in_code_block = False

    def emit_paragraph() -> None:
        nonlocal paragraph, title_stack
        block = "\n".join(paragraph).strip()
        paragraph = []
        if not block:
            return
        heading = re.match(r"^(#{1,6})\s+(.+)$", block)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).strip()
            title_stack = title_stack[: level - 1]
            title_stack.append(title)
        blocks.append({"text": block, "title_path": list(title_stack), "type": "text"})

    def emit_code_block() -> None:
        nonlocal code_block
        block = "\n".join(code_block).strip()
        code_block = []
        if block:
            blocks.append({"text": block, "title_path": list(title_stack), "type": "code"})

    for line in str(markdown or "").splitlines():
        if line.lstrip().startswith("```"):
            if in_code_block:
                code_block.append(line)
                emit_code_block()
                in_code_block = False
            else:
                emit_paragraph()
                code_block = [line]
                in_code_block = True
            continue
        if in_code_block:
            code_block.append(line)
            continue
        if not line.strip():
            emit_paragraph()
            continue
        paragraph.append(line)

    if in_code_block:
        emit_code_block()
    emit_paragraph()
    return blocks


def _make_chunk(index: int, parts: list[str], title_path: list[str], *, content_type: str = "text") -> dict[str, Any]:
    text = "\n\n".join(part for part in parts if part).strip()
    return {
        "chunkId": f"raw_{index + 1:06d}",
        "ordinal": index,
        "text": text,
        "contextualText": "\n".join(title_path + [text]) if title_path else text,
        "titlePath": title_path,
        "contentType": content_type,
        "chunkStage": "raw",
        "sourceChunkIds": [],
        "metadata": {"parser": "fallback", "chunker": "markdown"},
    }


def _tail_text_by_lines(text: str, limit: int) -> str:
    if limit <= 0 or not text:
        return ""
    lines = [line for line in text.splitlines() if line.strip()]
    selected: list[str] = []
    total = 0
    for line in reversed(lines):
        line_len = len(line) + (1 if selected else 0)
        if selected and total + line_len > limit:
            break
        selected.append(line)
        total += line_len
        if total >= limit:
            break
    return "\n".join(reversed(selected)).strip()


def _write_text_artifact(storage: Any, storage_key: str, text: str) -> None:
    with tempfile.NamedTemporaryFile("w+b") as temp:
        temp.write(text.encode("utf-8"))
        temp.seek(0)
        storage.put_file(temp, storage_key)


def _write_json_artifact(storage: Any, storage_key: str, payload: dict[str, Any]) -> None:
    content = json.dumps(payload, ensure_ascii=False, indent=2)
    _write_text_artifact(storage, storage_key, content)


def post_parse_failure_callback(job_id: str, request: ParseDocumentJobRequest, error: str) -> None:
    post_callback(
        request.callback.url,
        request.callback.token,
        {
            "jobId": job_id,
            "status": "failed",
            "markdownKey": "",
            "jsonKey": "",
            "rawChunksKey": "",
            "rawChunkCount": 0,
            "ragChunksKey": "",
            "ragChunkCount": 0,
            "chunksKey": "",
            "chunkCount": 0,
            "error": error[:2000],
        },
    )
